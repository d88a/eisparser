"""Utilities for extracting text from different document formats."""

from __future__ import annotations

import logging
import os
import shutil
import tempfile
import zipfile
from pathlib import Path

import olefile
import openpyxl
import pdfplumber
import xlrd
from docx import Document

logger = logging.getLogger("text_extraction")
_word_com_unavailable_logged = False
_word_fallback_logged = False


def extract_text_from_pdf(path: str) -> str:
    chunks: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                chunks.append(page.extract_text() or "")
    except Exception as exc:
        return f"Error extracting text from PDF: {exc}"
    return "\n".join(chunks)


def extract_text_from_doc(path: str) -> str:
    global _word_fallback_logged

    converted = _extract_doc_via_word(path)
    if converted:
        return converted

    if not _word_fallback_logged:
        logger.info("Word DOC conversion fallback is active; DOC files will be parsed without COM.")
        _word_fallback_logged = True

    try:
        with open(path, "rb") as f:
            data = f.read()

        try:
            text = data.decode("utf-8", errors="ignore")
        except Exception:
            text = data.decode("cp1251", errors="ignore")

        cleaned = []
        for ch in text:
            if ch in ("\n", "\r", "\t") or ch.isprintable():
                cleaned.append(ch)
        result = "".join(cleaned)
        return result if result.strip() else "Error extracting text from DOC: empty content"
    except Exception as exc:
        return f"Error extracting text from DOC: {exc}"


def extract_text_from_docx(path: str) -> str:
    try:
        doc = Document(path)
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        result = "\n".join(parts)
        return result if result.strip() else "Error extracting text from DOCX: empty content"
    except Exception as exc:
        return f"Error extracting text from DOCX: {exc}"


def extract_text_from_excel(path: str) -> str:
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
        parts: list[str] = []
        for sheet in wb.worksheets:
            parts.append(f"=== Sheet {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                values = [str(v) for v in row if v is not None]
                if values:
                    parts.append(" ".join(values))
        result = "\n".join(parts)
        return result if result.strip() else "Error extracting text from Excel: empty content"
    except Exception as exc:
        return f"Error extracting text from Excel: {exc}"


def _extract_doc_via_word(path: str) -> str | None:
    global _word_com_unavailable_logged

    try:
        import win32com.client as win32  # type: ignore
    except ImportError:
        if not _word_com_unavailable_logged:
            logger.warning("Word COM is unavailable (win32com is missing). Falling back without DOC conversion.")
            _word_com_unavailable_logged = True
        return None

    tmp_dir = tempfile.mkdtemp(prefix="doc-convert-")
    word = None
    doc = None
    try:
        tmp_docx = Path(tmp_dir) / "converted.docx"
        try:
            word = win32.DispatchEx("Word.Application")
        except Exception as exc:
            if not _word_com_unavailable_logged:
                logger.warning("Word COM is unavailable (%s). Falling back without DOC conversion.", exc)
                _word_com_unavailable_logged = True
            return None

        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(path))
        doc.SaveAs(str(tmp_docx), FileFormat=16)
        doc.Close()
        doc = None

        return extract_text_from_docx(str(tmp_docx))
    except Exception as exc:
        logger.info("Word COM conversion failed for %s: %s", path, exc)
        return None
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()
        shutil.rmtree(tmp_dir, ignore_errors=True)


def extract_text_from_xls(path: str) -> str:
    try:
        wb = xlrd.open_workbook(path, on_demand=True)
        parts: list[str] = []
        for sheet in wb.sheets():
            parts.append(f"=== Sheet {sheet.name} ===")
            for rx in range(sheet.nrows):
                values = []
                for cx in range(sheet.ncols):
                    value = sheet.cell_value(rx, cx)
                    if value in ("", None):
                        continue
                    if isinstance(value, float) and value.is_integer():
                        values.append(str(int(value)))
                    else:
                        values.append(str(value))
                if values:
                    parts.append("\t".join(values))
        result = "\n".join(parts)
        return result if result.strip() else "Error extracting text from XLS: empty content"
    except Exception as exc:
        return f"Error extracting text from XLS: {exc}"


def _is_ole_excel(path: str) -> bool:
    try:
        if not olefile.isOleFile(path):
            return False
        with olefile.OleFileIO(path) as ole:
            entries = {" / ".join(entry) for entry in ole.listdir()}
        return any("Workbook" in entry or "Book" in entry for entry in entries)
    except Exception:
        return False


def extract_text_from_zip(path: str) -> str:
    temp_dir = None
    try:
        text_chunks: list[str] = []
        temp_dir = tempfile.mkdtemp(prefix="zip-extract-")
        with zipfile.ZipFile(path, "r") as zf:
            zf.extractall(temp_dir)

        for root, _dirs, files in os.walk(temp_dir):
            for fname in files:
                inner_path = os.path.join(root, fname)
                text = extract_text_from_any_file(inner_path)
                if text and not text.startswith("Error extracting"):
                    text_chunks.append(f"=== File inside archive: {fname} ===\n{text}")
        return "\n\n".join(text_chunks)
    except Exception as exc:
        return f"Error extracting text from ZIP: {exc}"
    finally:
        if temp_dir and os.path.isdir(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)


def detect_type_by_extension(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext in (".xlsx", ".xlsm"):
        return "xlsx"
    if ext == ".xls":
        return "xls"
    if ext == ".doc":
        return "doc"
    if ext == ".zip":
        return "zip"
    return "unknown"


def detect_type_by_signature(path: str) -> str:
    try:
        with open(path, "rb") as f:
            header = f.read(2048)
    except Exception:
        return "unknown"

    if header.startswith(b"%PDF"):
        return "pdf"
    if header.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
        if _is_ole_excel(path):
            return "xls"
        return "doc"
    if header.startswith((b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")):
        try:
            with zipfile.ZipFile(path, "r") as zf:
                names = zf.namelist()
        except Exception:
            return "zip"
        if any(name.startswith("word/") or name.endswith(".docx") for name in names):
            return "docx"
        if any(name.startswith("xl/") or name.endswith(".xlsx") for name in names):
            return "xlsx"
        if any(name.endswith(".pdf") for name in names):
            return "pdf"
        return "zip"
    return "unknown"


def extract_text_from_any_file(path: str) -> str:
    ftype = detect_type_by_extension(path)
    signature_type = detect_type_by_signature(path)
    if signature_type != "unknown":
        ftype = signature_type

    if ftype == "pdf":
        return extract_text_from_pdf(path)
    if ftype == "docx":
        return extract_text_from_docx(path)
    if ftype == "xlsx":
        return extract_text_from_excel(path)
    if ftype == "xls":
        return extract_text_from_xls(path)
    if ftype == "doc":
        return extract_text_from_doc(path)
    if ftype == "zip":
        return extract_text_from_zip(path)
    return f"Unknown file type: {os.path.basename(path)}"
